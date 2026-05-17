#!/usr/bin/env python3
"""Generate project documentation in DOCX format for 妖妖乐 (Monster Mood Report)."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ---------------------------------------------------------------------------
#  Global style tweaks
# ---------------------------------------------------------------------------
style = doc.styles["Normal"]
font = style.font
font.name = "Microsoft YaHei"
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x5A, 0x2D, 0x1F)
    hs.font.bold = True
    hs.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

# ---------------------------------------------------------------------------
#  Helpers
# ---------------------------------------------------------------------------
def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Shading Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            tbl.rows[ri + 1].cells[ci].text = str(val)
    return tbl


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p


# ===================================================================
#  TITLE PAGE
# ===================================================================
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("妖妖乐\nMonster Mood Report")
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0xC0, 0x84, 0xFC)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("项目技术说明文档")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0xEC, 0x48, 0x99)

doc.add_paragraph()
ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ver.add_run(f"v0.3.0 — {datetime.date.today().isoformat()}")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ===================================================================
#  1. 项目概述
# ===================================================================
doc.add_heading("1  项目概述", level=1)

doc.add_paragraph(
    "妖妖乐（Monster Mood Report）是一个可爱风格的 AI 情绪人格报告生成应用。"
    "用户上传一段视频，系统通过视觉大模型分析视频内容，结合 DeepSeek 文本大模型，"
    "生成一只专属的「情绪小妖怪」——包含 MBTI 人格配方、能量值、情绪描述，"
    "并支持进入「妖妖占卜屋」与五只情绪小妖怪互动问答，最终生成可保存/分享的结果卡片。"
)

doc.add_heading("1.1  核心功能", level=2)
features = [
    "🎬 视频上传与分析：支持 mp4/mov/webm 格式，通过 vLLM + Qwen3-VL-32B 多模态模型理解画面内容",
    "🦄 专属妖怪生成：基于视频分析结果，LLM 动态生成独一无二的妖怪名字、属性、外表描述",
    "🔮 MBTI 人格配方：生成 4 维 MBTI 混合人格报告（含百分比和可爱文案）",
    "⚡ 今日能量值：0-100 的动态能量评分",
    "📊 浏览行为分析：模拟用户画像和行为的趣味分析",
    "🏠 妖妖占卜屋：五只固定情绪小妖怪（乐啵啵/灰绵绵/怯团团/嫌叽叽/炸毛毛）为用户答疑解惑",
    "🎴 分享卡片：最终生成可保存/分享的梦幻风格结果卡片",
]
for f in features:
    add_bullet(doc, f)

doc.add_heading("1.2  目标用户", level=2)
doc.add_paragraph(
    "面向年轻用户群体（18-30 岁），尤其是对 MBTI、情绪分析、心理测试感兴趣的用户。"
    "产品以轻松娱乐的方式提供「人格配方」体验，不涉及真实心理评估。产品定位为轻娱乐 Demo。"
)

# ===================================================================
#  2. 技术栈
# ===================================================================
doc.add_heading("2  技术栈", level=1)

doc.add_heading("2.1  前端", level=2)
frontend_data = [
    ("框架", "React 19 + TypeScript"),
    ("构建工具", "Vite 7 + @tanstack/react-start"),
    ("路由", "@tanstack/react-router (文件路由)"),
    ("UI 组件库", "Radix UI + shadcn/ui"),
    ("动画引擎", "Framer Motion 12"),
    ("样式方案", "Tailwind CSS 4 + tw-animate-css"),
    ("数据可视化", "Recharts"),
    ("HTTP 客户端", "原生 fetch + FormData"),
    ("状态管理", "局部 useState + 模块级可变对象 (yaoyao-data.ts)"),
]
add_table(doc, ["类别", "技术选型"], frontend_data)

doc.add_heading("2.2  后端", level=2)
backend_data = [
    ("Web 框架", "Python FastAPI"),
    ("异步运行时", "uvicorn + asyncio"),
    ("数据验证", "Pydantic v2"),
    ("HTTP 客户端", "httpx (异步)"),
    ("视频帧提取", "ffmpeg + ffprobe"),
    ("进程管理", "PM2 (ecosystem.config.cjs)"),
]
add_table(doc, ["类别", "技术选型"], backend_data)

doc.add_heading("2.3  AI 模型", level=2)
ai_data = [
    ("视觉模型", "Qwen3-VL-32B-Instruct-FP8"),
    ("推理引擎", "vLLM 0.19.1 (2× RTX 5090, Tensor Parallel)"),
    ("文本模型", "DeepSeek Chat (deepseek-chat / deepseek-v3.2)"),
    ("备用视觉", "豆包 doubao-seed-2-0-lite (OpenAI-Next 中转)"),
    ("模型部署", "vLLM Server :8100"),
]
add_table(doc, ["类别", "技术选型"], ai_data)

doc.add_heading("2.4  部署运维", level=2)
ops_data = [
    ("进程管理器", "PM2 (开机自启、崩溃重启)"),
    ("GPU 驱动", "NVIDIA 575.64.03 (CUDA 12.8)"),
    ("Python 环境", "Miniconda Python 3.13 (CUDA 12.8)"),
    ("Node 版本", "Node.js 24+"),
    ("操作系统", "Linux (Ubuntu 22.04)"),
]
add_table(doc, ["类别", "技术选型"], ops_data)

# ===================================================================
#  3. 项目结构
# ===================================================================
doc.add_heading("3  项目结构", level=1)

add_code(doc, """monster-mood-report/
├── src/                          # 前端源代码
│   ├── components/
│   │   ├── yaoyao/
│   │   │   ├── YaoyaoApp.tsx     # 7 步流程主组件 (upload→loading→monster→report→
│   │   │   │                     #   transition→questions→answers→card)
│   │   │   ├── CuteMonster.tsx   # UI 骨架组件 (ScreenFrame + 复古怪兽动画)
│   │   │   └── YaoyaoApp.tsx     # 所有 8 个 Screen 组件（见 4.2 节）
│   │   ├── AnimatedMonster.tsx   # 可定制 SVG 动画妖怪组件
│   │   └── ui/                   # shadcn/ui 组件库（~50 个原子组件）
│   ├── lib/
│   │   ├── api.ts                # 后端 API 调用封装 (uploadVideo / pollTask / generateAnswers)
│   │   ├── yaoyao-data.ts        # 全局数据存储 + 默认 demo 数据
│   │   ├── utils.ts              # 通用工具函数
│   │   ├── error-capture.ts      # 错误捕获
│   │   └── error-page.tsx        # 错误页面
│   ├── routes/index.tsx          # 首页路由
│   ├── router.tsx                # TanStack Router 配置
│   ├── routeTree.gen.ts          # 自动生成的路由树
│   ├── styles.css                # 全局样式 + CSS 动画
│   ├── server.ts                 # TanStack Start 服务端入口
│   └── start.ts                  # 启动入口
├── backend/                      # 后端服务
│   ├── main.py                   # FastAPI 应用 + 4 个路由 + 异步任务管线
│   ├── services.py               # LLM/Vision 调用层 + JSON 解析
│   ├── prompts.py                # 4 个 Prompt 模板（分析/元数据/报告/回答）
│   ├── models.py                 # Pydantic 数据模型
│   ├── config.py                 # Provider 配置 + API Key
│   └── start.sh                  # shell 启动脚本
├── public/                       # 静态资源
│   ├── monster.mp4 / monster.png # 中心妖怪动画
│   ├── 黑夜转身.mp4              # 转场动画
│   ├── 乐啵啵.mp4 / 灰绵绵.mp4   # 情绪小妖怪 mp4
│   ├── 怯团团.mp4 / 嫌叽叽.mp4   #
│   ├── 炸毛毛.mp4                #
│   ├── 妖妖占卜屋.png            # QuestionsScreen 背景
│   └── 妖妖占卜屋screen5.png     # AnswersScreen 背景
├── ecosystem.config.cjs          # PM2 进程配置 (3 apps)
├── package.json                  # Node 依赖
├── vite.config.ts                # Vite 构建配置
├── tsconfig.json                 # TypeScript 配置
└── components.json               # shadcn/ui 配置""")

# ===================================================================
#  4. 前端架构详解
# ===================================================================
doc.add_heading("4  前端架构详解", level=1)

doc.add_heading("4.1  技术选型说明", level=2)
doc.add_paragraph(
    "前端基于 TanStack Start（React 19 + Vite 7 + 文件路由）构建。UI 层使用 Radix UI 无头组件"
    "+ shadcn/ui 的封装组件，配合 Tailwind CSS 4 实现原子化样式。动画部分全部由 Framer Motion 12 驱动。\n\n"
    "数据状态管理采用了极简方案——模块级可变对象（yaoyao-data.ts），而非 Redux/Zustand 等状态库。"
    "后端返回的 ReportData 通过 Object.assign 直接修改同一引用，无需改动任何 Screen 组件代码即可"
    "从 demo 数据切换到真实数据。"
)

doc.add_heading("4.2  页面流程（7 步）", level=2)
doc.add_paragraph("整个应用是一个线性七步流，由 YaoyaoApp.tsx 中的 step state 控制：")
flow_data = [
    ("UploadScreen", "文件选择与上传", "支持 mp4/mov/webm，最大 200MB，包含视频预览"),
    ("LoadingScreen", "后端处理等待", "旋转水晶球动画 + 浮动星星 + 渐进提示文案"),
    ("MonsterScreen", "专属妖怪展示", "竖版海报布局：妖怪名字 + SVG 动画 + 诗句卡片 + 标签"),
    ("ReportScreen", "今日状态报告", "MBTI 叶子花瓣图 + 旋转轨道 + 能量条 + 情绪云朵"),
    ("TransitionScreen", "转场动画", "全屏星空动画，「黑夜转身.mp4」飞入 + 占卜屋浮现"),
    ("QuestionsScreen", "妖妖占卜屋", "梦幻全屏背景 + 3 个渐变磨砂问题卡片"),
    ("AnswersScreen", "情绪妖怪回答", "5 只情绪 mp4 圆环布局 → 选中后显示回答解释卡"),
    ("CardScreen", "分享卡片", "梦幻紫粉渐变卡片 + MBTI/能量/问答 + 保存/分享/再来一次"),
]
add_table(doc, ["页面", "功能", "说明"], flow_data)

doc.add_heading("4.3  数据流", level=2)
doc.add_paragraph("数据流分为三个阶段：")
add_bullet(doc, "上传阶段：用户选择文件 → UploadScreen 调 api.uploadVideo() → POST /api/upload-video → 返回 task_id")
add_bullet(doc, "轮询阶段：LoadingScreen 调 api.pollTask() → 每 2 秒 GET /api/task/{id}/status → 拿到 ReportData")
add_bullet(doc, "回答阶段：用户选择问题 → generateAnswers() → POST /api/generate-answers → 更新 emotionMonsters.answer")

doc.add_paragraph(
    "三个阶段均通过模块级函数 setYaoyaoData() 写回全局 yaoyaoData 对象，"
    "Screen 组件在 step state 切换时自动读到最新数据，无需额外 props 传递。"
)

doc.add_heading("4.4  组件树", level=2)
add_code(doc, """YaoyaoApp (step state machine)
├── Header                    # 标题栏 + 品牌标识
├── UploadScreen              # 文件选择/预览/上传
├── LoadingScreen             # 等待动画 + 轮询进度
├── MonsterScreen             # 专属妖怪展示
│   └── AnimatedMonster       # SVG 动画妖怪
├── ReportScreen              # 今日状态报告
│   ├── MbtiPetal ×4          # MBTI 叶子组件
│   ├── CenterMonster         # 中心 mp4 妖怪
│   └── Cloud ×2              # 能量值 + 情绪云朵
├── TransitionScreen          # 转场动画
├── QuestionsScreen           # 妖妖占卜屋 - 问题选择
├── AnswersScreen             # 情绪妖怪回答
│   └── MonsterBubble ×5      # 圆形 mp4 气泡
└── CardScreen                # 分享卡片
    ├── Star4                 # 四角星 SVG
    ├── CloudPuff             # 云朵 SVG
    ├── InfoTile ×2           # MBTI + 能量块
    └── ResultButton ×3       # 保存/分享/再玩一次""")

# ===================================================================
#  5. 后端架构详解
# ===================================================================
doc.add_heading("5  后端架构详解", level=1)

doc.add_heading("5.1  路由与 API", level=2)
doc.add_paragraph("FastAPI 应用定义在 backend/main.py 中，提供 4 个路由端点：")
api_data = [
    ("POST", "/api/upload-video", "上传视频 → 创建任务 → 返回 task_id", "TaskResponse"),
    ("GET", "/api/task/{task_id}/status", "轮询任务状态 → 获取分析报告", "TaskResponse"),
    ("POST", "/api/generate-answers", "为选定问题生成五只妖怪的回答", "GenerateAnswersResponse"),
    ("GET", "/api/health", "健康检查", "JSON"),
]
add_table(doc, ["方法", "路径", "功能", "响应模型"], api_data)

doc.add_heading("5.2  异步任务管线", level=2)
doc.add_paragraph("上传视频后，后端立即返回 task_id，通过 asyncio.create_task 启动后台任务：")
add_code(doc, """1. upload_video()
   → 校验文件类型 & 大小
   → 生成 task_id, 保存文件
   → asyncio.create_task(_process_video(task_id, path))
   → 立即返回 { task_id, status: "processing" }

2. _process_video() [后台]
   ├── analyze_video(file_path)
   │   ├── 若 vllm:  ffmpeg 提取 4 帧 → base64 → vLLM API 分析
   │   ├── 若 doubao: 上传文件 → 豆包 API 分析
   │   └── 若 none:   ffprobe 元数据 → LLM 模拟分析
   ├── generate_report(video_analysis)
   │   └── DeepSeek Chat + json_object → 结构化报告 JSON
   ├── _normalize_report()
   │   ├── 注入固定 5 只情绪妖怪 + 清空 answer
   │   ├── 修正 MBTI 百分比到 100
   │   ├── 补齐 videoAnalysis 到 4 条
   │   └── 补齐 recommendedQuestions 到 3+
   └── tasks[task_id] = { status: "completed", result: ReportData }

3. generate_answers(video_analysis, question)
   └── DeepSeek Chat → 5 只妖怪回答 JSON""")

doc.add_heading("5.3  视觉模型集成 (vLLM)", level=2)
doc.add_paragraph(
    "视频分析默认使用 vLLM 部署的 Qwen3-VL-32B-Instruct-FP8 多模态模型。流程如下：\n\n"
    "1. ffprobe 获取视频时长\n"
    "2. ffmpeg 按 fps 均匀提取最多 4 帧（下采样至 480px 宽）\n"
    "3. 每帧转为 base64 data-URL\n"
    "4. 构造 OpenAI-compatible 的 chat/completions 请求（type: image_url）\n"
    "5. 模型返回视频内容分析文本\n\n"
    "vLLM 服务部署在 localhost:8100，PM2 进程管理，开机自启。"
    "模型以 FP8 精度加载，Tenso Parallel=2，占用 ~44GB（2×RTX 5090 各 ~22GB）。"
)

doc.add_heading("5.4  文本模型集成 (DeepSeek)", level=2)
doc.add_paragraph(
    "报告生成和妖妖怪回答均通过 DeepSeek Chat API (deepseek-chat) 完成。"
    "使用 response_format: json_object 确保输出为 JSON。"
    "支持热切换至 OpenAI-Next 兼容 API（通过 config.py 的 LLM_PROVIDER 环境变量）。\n\n"
    "JSON 解析兼容处理：后端配置 json.loads(strict=False)，允许 DeepSeek 生成的 JSON 中含有未转义的控制字符（换行符等）。"
)

doc.add_heading("5.5  数据模型", level=2)
add_code(doc, """TaskStatus : PROCESSING / COMPLETED / FAILED

ReportData:
├── monster: Monster
│   ├── name: str          # 妖怪名字（3字）
│   ├── type: str          # 妖怪类型
│   ├── emoji: str         # 一个 emoji
│   ├── color: str         # 主题色 hex
│   ├── attributes[]       # [{label, value}] 4个属性
│   └── intro: str         # 引导语
├── mbtiMix[]              # [{type, percent, color, cute}] 4个
├── energyScore: int       # 0-100
├── emotionText: str       # 今日情绪描述
├── videoAnalysis[]        # [{icon, text}] 4条
├── recommendedQuestions[] # 4个问题
└── emotionMonsters[]      # 5只固定情绪妖怪""")

# ===================================================================
#  6. Prompt 工程
# ===================================================================
doc.add_heading("6  Prompt 工程", level=1)

doc.add_paragraph("系统包含 4 个精心设计的 Prompt 模板，位于 backend/prompts.py：")

prompt_data = [
    ("VIDEO_ANALYSIS_PROMPT", "视觉模型用", "分析视频主题、情绪基调、观众行为和画像。要求轻松有趣，禁严肃心理学术语"),
    ("VIDEO_METADATA_PROMPT", "纯文本备选", "当无视觉模型时，根据元数据想象推测。禁止提及「推测」「可能」等弱化语"),
    ("REPORT_SYSTEM_PROMPT", "报告生成", "生成结构化 JSON 报告。含 10 个子字段的精确格式约束。强调萌系治愈风格"),
    ("ANSWER_SYSTEM_PROMPT", "妖怪回答", "五只妖怪各 30-80 字回答。严格定义每只性格特征和回答风格"),
]
add_table(doc, ["Prompt 名称", "使用场景", "关键约束"], prompt_data)

doc.add_paragraph(
    "Prompt 设计特点：\n"
    "• 角色定义：数据炼妖师 / 情绪小妖怪——将技术输出包装为趣味叙事\n"
    "• 拒绝列表：明确禁止「诊断」「症状」「障碍」等严肃词汇\n"
    "• 格式精确：JSON 的字段名、类型、长度、数量均有约束\n"
    "• 创造力要求：每次生成的妖怪名字、属性不能千篇一律\n"
    "• 回答风格：五只妖怪性格固定但回答内容需针对用户问题"
)

# ===================================================================
#  7. 部署与运维
# ===================================================================
doc.add_heading("7  部署与运维", level=1)

doc.add_heading("7.1  PM2 进程管理", level=2)
doc.add_paragraph("所有服务通过 PM2 统一管理，配置在 ecosystem.config.cjs 中：")
pm2_data = [
    ("yaoyao-backend", "uvicorn main:app", "8787", "Python FastAPI"),
    ("yaoyao-frontend", "vite dev --host", "5173", "Vite Dev Server"),
    ("vllm-server", "vllm.entrypoints.openai.api_server", "8100", "Qwen3-VL-32B"),
]
add_table(doc, ["服务名称", "启动命令", "端口", "说明"], pm2_data)

doc.add_paragraph("PM2 配置了崩溃自动重启（max_restarts=5-10），以及 3-15 秒的重启间隔。")

doc.add_heading("7.2  环境变量", level=2)
env_data = [
    ("YAOYAO_LLM_PROVIDER", "deepseek (默认) / openai-next", "文本 LLM 提供者"),
    ("YAOYAO_VISION_PROVIDER", "vllm (默认) / doubao", "视觉模型提供者"),
    ("DEEPSEEK_API_KEY", "DeepSeek API Key", "DeepSeek 鉴权"),
    ("YAOYAO_API_KEY", "OpenAI-Next API Key", "OpenAI 兼容 API 鉴权"),
]
add_table(doc, ["变量名", "可选值", "说明"], env_data)

doc.add_heading("7.3  硬件要求", level=2)
hw_data = [
    ("GPU", "2× NVIDIA RTX 5090 (32GB+ VRAM 每张)"),
    ("显存占用", "~22GB/GPU (Qwen3-VL-32B-FP8)"),
    ("系统内存", "推荐 64GB+"),
    ("磁盘", "SSD, 50GB+"),
    ("CUDA Driver", "≥ 575.64.03 (支持 CUDA 12.8)"),
]
add_table(doc, ["资源", "要求"], hw_data)

doc.add_heading("7.4  启动与维护命令", level=2)
add_code(doc, """# 启动所有服务
pm2 start ecosystem.config.cjs

# 重启后端 (代码更新后)
pm2 restart yaoyao-backend

# 查看日志
pm2 logs yaoyao-backend

# 查看 vLLM 模型加载状态
curl http://localhost:8100/v1/models

# 健康检查
curl http://localhost:8787/api/health

# 保存进程列表 (开机自启)
pm2 save --force""")

# ===================================================================
#  8. 开发指南
# ===================================================================
doc.add_heading("8  开发指南", level=1)

doc.add_heading("8.1  本地环境搭建", level=2)
add_code(doc, """# 前端
npm install
npm run dev     # → http://localhost:5173

# 后端
cd backend
pip install -r requirements.txt
python3 -m uvicorn main:app --reload --host 0.0.0.0 --port 8787

# vLLM 服务 (需要 2×RTX 5090)
python3 -m vllm.entrypoints.openai.api_server \\
    --model /root/Qwen3-VL-32B-Instruct-FP8 \\
    --port 8100 --tensor-parallel-size 2 \\
    --gpu-memory-utilization 0.90 --dtype bfloat16""")

doc.add_heading("8.2  常见开发任务", level=2)

doc.add_heading("修改 Prompt", level=3)
doc.add_paragraph("编辑 backend/prompts.py → pm2 restart yaoyao-backend。提示词修改无需重启 vLLM。")

doc.add_heading("更换视觉模型", level=3)
doc.add_paragraph("设置环境变量 YAOYAO_VISION_PROVIDER=doubao 或修改 config.py 中的 VISION_PROVIDER。")

doc.add_heading("更换文本模型", level=3)
doc.add_paragraph("设置 YAOYAO_LLM_PROVIDER=openai-next，并在 config.py 中配置对应的 API Key 和模型名。")

doc.add_heading("修改前端 UI", level=3)
doc.add_paragraph("所有 Screen 组件集中在 src/components/yaoyao/YaoyaoApp.tsx。样式使用 Tailwind CSS + 内联 style。色彩常量（BRAND_DARK 等）在文件头部定义。")

doc.add_heading("添加新页面", level=3)
doc.add_paragraph("在 YaoyaoApp.tsx 的 Step 联合类型中添加新步骤名，在对应位置渲染新 Screen 组件，在 go() 调用链中插入跳转。")

# ===================================================================
#  9. 故障排查
# ===================================================================
doc.add_heading("9  故障排查", level=1)

trouble_data = [
    ("视频分析失败", "检查 vLLM 服务状态: curl :8100/v1/models\n检查日志: pm2 logs vllm-server", "GPU 显存不足 / vLLM 未完全加载 / 视频格式不支持"),
    ("JSON 解析报错", "后端使用 json.loads(strict=False)，若仍有问题可在前端也加 try-catch", "LLM 输出了不合法 JSON / 控制字符未转义"),
    ("CUDA 驱动错误", "使用 miniconda Python 3.13（CUDA 12.8），而非系统 /usr/bin/python3（CUDA 13）", "NVIDIA 驱动版本与 PyTorch CUDA 版本不匹配"),
    ("前端白屏/NPM 错误", "npm install 重新安装依赖 / 检查 Node 版本 ≥ 24", "依赖版本冲突 / Node 版本过低"),
    ("端口冲突", "lsof -i :<port> 查找占用进程", "其他服务占用端口"),
    ("LLM API 超时", "检查 API Key 是否有效 / 网络是否可达 api.deepseek.com", "网络问题 / API Key 过期 / 余额不足"),
]
add_table(doc, ["问题", "排查方法", "常见原因"], trouble_data)

# ===================================================================
#  10. 版本历史
# ===================================================================
doc.add_heading("10  版本历史", level=1)
ver_data = [
    ("0.1.0", "初始版本", "前端 6 步 UI + demo 数据"),
    ("0.2.0", "后端集成", "FastAPI + vLLM + DeepSeek 完整管线"),
    ("0.3.0", "生产部署", "PM2 管理 + 错误处理 + CUDA 兼容性修复"),
]
add_table(doc, ["版本", "日期", "变更内容"], ver_data)

# ===================================================================
#  SAVE
# ===================================================================
output_path = "/root/monster-mood-report/妖妖乐_项目技术说明文档.docx"
doc.save(output_path)
print(f"✅ Document saved to: {output_path}")
